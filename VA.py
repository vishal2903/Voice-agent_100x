import asyncio
import logging
import os
import re
from typing import Optional
from datetime import datetime

from dotenv import load_dotenv
from openai import OpenAI

from livekit import agents
from livekit.agents import AgentSession, Agent, RoomInputOptions, function_tool, RunContext
from livekit.plugins import openai as lk_openai, noise_cancellation, silero

# For saving files
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

import mimetypes

# Google Drive
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request

# --- NEW IMPORTS for Gmail ---
import base64
from email.mime.text import MIMEText
from googleapiclient.discovery import build as gbuild
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request as GRequest
from google.oauth2.credentials import Credentials as GCredentials

load_dotenv()


GOOGLE_CREDENTIALS = os.getenv("GOOGLE_DRIVE_CREDENTIALS", r"C:\Users\visha\Downloads\Voice Agent\credentials.json")
GOOGLE_TOKEN       = os.getenv("GOOGLE_DRIVE_TOKEN",       r"C:\Users\visha\Downloads\Voice Agent\token.json")
GOOGLE_FOLDER_ID   = os.getenv("GOOGLE_DRIVE_FOLDER_ID",   "")  # optional
DRIVE_SCOPES = ["https://www.googleapis.com/auth/drive.file"]  # safer scope

# --- NEW GMAIL CONSTANTS  ---
GMAIL_SCOPES = ["https://www.googleapis.com/auth/gmail.send"]
GOOGLE_OAUTH_CLIENT = os.getenv("GOOGLE_OAUTH_CLIENT", r"C:\Users\visha\Downloads\Voice Agent\credentials.json")
GMAIL_TOKEN_PATH = os.getenv("GMAIL_TOKEN_PATH", r"C:\Users\visha\Downloads\Voice Agent\gmail_token.json")

# ---------- Config ----------
# Vector store for File Search (Responses API)
VECTOR_STORE_ID = os.getenv("VECTOR_STORE_ID")

# Primary & fallback save locations
SAVE_DIR_PRIMARY = os.getenv("SAVE_DIR", r"C:\Users\visha\Downloads")
SAVE_DIR_FALLBACK = r"C:\temp"

# Logging noise reduction
logging.getLogger("livekit.agents").setLevel(logging.WARNING)
logging.getLogger("livekit.plugins.silero").setLevel(logging.ERROR)
#logging.info("Sending email → to=%s, subject=%s", self._pending_email["to"], self._pending_email["subject"])


# ---------- Helpers ----------
def _slugify(s: str) -> str:
    s = (s or "report").lower()
    s = re.sub(r"[^a-z0-9]+", "-", s).strip("-")
    return s[:60] or "report"

def _ensure_dir(path: str):
    os.makedirs(path, exist_ok=True)

def _save_as_docx(path: str, topic: str, content: str):
    doc = Document()
    doc.add_heading(f"Research Report: {topic}", 0)
    for para in (content or "").split("\n\n"):
        doc.add_paragraph(para)
    doc.save(path)

def _save_as_pdf(path: str, topic: str, content: str):
    c = canvas.Canvas(path, pagesize=letter)
    width, height = letter
    margin_x, margin_y = 72, 72  # 1 inch margins

    # Title
    c.setFont("Helvetica-Bold", 14)
    c.drawString(margin_x, height - margin_y, f"Research Report: {topic}")

    # Body
    c.setFont("Helvetica", 11)
    y = height - margin_y - 24
    max_chars = 95

    def draw_line(line: str, y_pos: float):
        c.drawString(margin_x, y_pos, line)

    for raw_line in (content or "").splitlines():
        line = raw_line
        if not line:
            y -= 14
            if y < margin_y:
                c.showPage()
                c.setFont("Helvetica", 11)
                y = height - margin_y
            continue

        # naive wrapping
        while len(line) > max_chars:
            draw_line(line[:max_chars], y)
            line = line[max_chars:]
            y -= 14
            if y < margin_y:
                c.showPage()
                c.setFont("Helvetica", 11)
                y = height - margin_y
        draw_line(line, y)
        y -= 14
        if y < margin_y:
            c.showPage()
            c.setFont("Helvetica", 11)
            y = height - margin_y

    c.save()

def _save_report_to_file(topic: str, content: str, file_type: str) -> str:
    """Try primary dir; fall back to C:\\temp. Returns the final path."""
    if file_type not in ("docx", "pdf"):
        raise ValueError("file_type must be 'docx' or 'pdf'")

    timestamp = datetime.now().strftime("%Y-%m-%d-%H-%M-%S")
    base = f"research_report-{_slugify(topic)}-{timestamp}"
    filename = base + (".docx" if file_type == "docx" else ".pdf")

    # Try primary
    try:
        _ensure_dir(SAVE_DIR_PRIMARY)
        path = os.path.join(SAVE_DIR_PRIMARY, filename)
        if file_type == "docx":
            _save_as_docx(path, topic, content)
        else:
            _save_as_pdf(path, topic, content)
        return path
    except Exception as e:
        logging.warning("Save failed in %s: %s. Falling back to %s",
                        SAVE_DIR_PRIMARY, e, SAVE_DIR_FALLBACK)

    # Fallback
    _ensure_dir(SAVE_DIR_FALLBACK)
    path = os.path.join(SAVE_DIR_FALLBACK, filename)
    if file_type == "docx":
        _save_as_docx(path, topic, content)
    else:
        _save_as_pdf(path, topic, content)
    return path


# ---------- Drive integration ----------
def _ensure_drive_service():
    """Return an authorized Drive v3 service; performs first-time OAuth if needed."""
    creds = None
    if os.path.exists(GOOGLE_TOKEN):
        creds = Credentials.from_authorized_user_file(GOOGLE_TOKEN, DRIVE_SCOPES)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            if not os.path.exists(GOOGLE_CREDENTIALS):
                raise FileNotFoundError(f"Missing Google credentials JSON at {GOOGLE_CREDENTIALS}")
            flow = InstalledAppFlow.from_client_secrets_file(GOOGLE_CREDENTIALS, DRIVE_SCOPES)
            # Desktop flow (opens browser). If port collision, try again with a different port.
            try:
                creds = flow.run_local_server(port=0)
            except Exception:
                creds = flow.run_console()
        with open(GOOGLE_TOKEN, "w", encoding="utf-8") as token:
            token.write(creds.to_json())

    return build("drive", "v3", credentials=creds)

def _guess_mime(file_path: str) -> str:
    # Explicit known types; fallback to mimetypes
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return "application/pdf"
    if ext == ".docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    mt, _ = mimetypes.guess_type(file_path)
    return mt or "application/octet-stream"

def upload_file_to_drive(file_path: str, folder_id: Optional[str] = None) -> dict:
    """Upload local file to Drive. Returns dict with id, webViewLink."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Not found: {file_path}")

    service = _ensure_drive_service()
    body = {"name": os.path.basename(file_path)}
    if folder_id:
        body["parents"] = [folder_id]

    media = MediaFileUpload(file_path, mimetype=_guess_mime(file_path), resumable=True)
    file = service.files().create(
        body=body,
        media_body=media,
        fields="id, name, webViewLink, webContentLink, parents"
    ).execute()
    # Request a webViewLink (owner can open immediately). Sharing is not changed.
    return file

#------- for GMAIL ------------

def _ensure_gmail_service():
    """
    Ensure an authenticated Gmail API service (uses a separate token file from Drive).
    First run will open a browser to consent and create GMAIL_TOKEN_PATH.
    """
    if not os.path.exists(GOOGLE_OAUTH_CLIENT):
        raise FileNotFoundError(f"Missing Google OAuth client JSON at {GOOGLE_OAUTH_CLIENT}")

    creds = None
    if os.path.exists(GMAIL_TOKEN_PATH):
        creds = GCredentials.from_authorized_user_file(GMAIL_TOKEN_PATH, GMAIL_SCOPES)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(GRequest())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(GOOGLE_OAUTH_CLIENT, GMAIL_SCOPES)
            # Use local server flow so Windows pops a browser to consent once
            creds = flow.run_local_server(port=0)
        with open(GMAIL_TOKEN_PATH, "w", encoding="utf-8") as token:
            token.write(creds.to_json())

    return gbuild("gmail", "v1", credentials=creds)

def send_email_via_gmail(to_email: str, subject: str, body: str) -> dict:
    """
    Build a simple text email and send it using Gmail API.
    Returns the Gmail API response (message id, etc).
    """
    service = _ensure_gmail_service()

    msg = MIMEText(body, _charset="utf-8")
    msg["to"] = to_email
    msg["subject"] = subject

    raw = base64.urlsafe_b64encode(msg.as_bytes()).decode("utf-8")
    sent = service.users().messages().send(userId="me", body={"raw": raw}).execute()
    return sent



# ---------- Agent ----------
class VoiceAssistant(Agent):
    def __init__(self) -> None:
        super().__init__(instructions="""You are a fast, efficient voice AI assistant for an AI instructor running a 6-month GenAI cohort.

CORE BEHAVIOR
- Keep answers concise, classroom-ready, and spoken clearly.
- Use FILE SEARCH for foundational concepts from internal notes (vector store).
- Use WEB SEARCH for recent developments and news (today/recent/latest).
- If the user says 'save', 'save as', 'export', or confirms saving, call the tool `save_last_report` with the requested file_type (default: docx). Do not say you cannot save.
- If the user says ‘upload to Drive’, call upload_last_report_to_drive (after saving). Perform at most two tools: save

EMAIL FLOW
- Always preview any email you draft and ask for explicit confirmation before sending it.
- After you show a draft, if the user replies with yes/okay/send it/go ahead, call `send_email` immediately. Do not call `compose_email` again unless the recipient or topic changed.

REPORT FORMATS (prompt-only guidance)
- daily_update: short dated bulle
- lesson_brief: definitions from files first; recent updates from web; demos & exercises.
- research_report: TL;DR, landscape, SOTA (dated), pitfalls, teaching recs, sources.

Always preview any email you draft and ask for explicit confirmation before sending it.
If the user confirms with ‘yes’, ‘send it’, ‘go ahead’, etc., call send_email immediately; do not draft again.

No filler; speak decisively.""")
        self.openai_client = OpenAI()

        # In-memory buffer for the last generated report
        self._last_report_topic: Optional[str] = None
        self._last_report_content: Optional[str] = None
        self._last_report_format: Optional[str] = None
        self._last_saved_path: Optional[str] = None
        self._pending_email = None  # NEW: holds the last composed email dict


    # ---- Web search (Responses API tool) ----
    @function_tool(
        description="Search the web for current/recent information and return direct facts."
    )
    async def web_search(self, context: RunContext, query: str) -> str:
        try:
            resp = self.openai_client.responses.create(
                model="gpt-4o-mini",
                tools=[{"type": "web_search"}],  # per docs
                input=query,
            )
            return resp.output_text or ""
        except Exception as e:
            logging.exception("web_search failed")
            return f"Search error: {e}"


    # ---- File search (Responses API tool + vector store) ----
    @function_tool(
        description="Search uploaded cohort/module notes using File Search (vector store)."
    )
    async def file_search(self, context: RunContext, query: str) -> str:
        if not VECTOR_STORE_ID:
            return "File search not configured. Set VECTOR_STORE_ID in your .env."
        try:
            resp = self.openai_client.responses.create(
                model="gpt-4o-mini",
                tools=[{
                    "type": "file_search",
                    "vector_store_ids": [VECTOR_STORE_ID],
                }],
                input=query,
            )

            return resp.output_text or ""
        except Exception as e:
            logging.exception("file_search failed")
            return f"File search error: {e}"


    # ---- Deep research (combine file_search + web_search) ----
    @function_tool(
        description=("Combine File Search (internal notes) and Web Search (recent) to produce a formatted brief. "
                     "Args: topic (str), format (daily_update|lesson_brief|research_report), "
                     "max_sources (int, default 6), recency_hint (str, default 'last 30 days').")
    )
    async def deep_research_report(
        self,
        context: RunContext,
        topic: str,
        format: str = "lesson_brief",
        max_sources: int = 6,
        recency_hint: str = "last 30 days",
    ) -> str:
        if not VECTOR_STORE_ID:
            return "File search not configured. Set VECTOR_STORE_ID in your .env."

        formats = {
            "daily_update": (
                "Return a DAILY UPDATE:\n"
                "1) TL;DR — 3 concise bullets with dates\n"
                f"2) What changed in {recency_hint} — 3–5 items, dated\n"
                f"3) Must-know links — up to {max_sources} URLs"
            ),
            "lesson_brief": (
                "Return a LESSON BRIEF for the next cohort session:\n"
                "1) TL;DR — what to teach & why (3 bullets)\n"
                "2) Core concepts & definitions — PULL FROM INTERNAL FILES FIRST\n"
                f"3) What's new on the web ({recency_hint}) — 3–5 dated points\n"
                "4) Demo ideas — one no-code and one code\n"
                "5) Exercises — 2–3 mini projects (~60–90 minutes)\n"
                f"6) Reading list — up to {max_sources} links (mix: internal + web)"
            ),
            "research_report": (
                "Return a RESEARCH REPORT:\n"
                "1) TL;DR — 5 bullets\n"
                "2) Landscape — key approaches & players\n"
                f"3) State of the art ({recency_hint}) — dated\n"
                "4) Risks / pitfalls\n"
                "5) Teaching recommendations — lecture flow for the cohort\n"
                f"6) Sources — up to {max_sources} URLs with dates"
            ),
        }
        section_spec = formats.get(format, formats["lesson_brief"])

        prompt = (
            "You assist an AI instructor planning a 6-month GenAI cohort.\n"
            "Use FILE SEARCH for foundational definitions/frameworks from internal notes.\n"
            f"Use WEB SEARCH for recent developments ({recency_hint}).\n"
            "Always include source URLs and dates when possible. Keep it concise and classroom-ready.\n\n"
            f"FORMAT:\n{section_spec}\n\n"
            f"TOPIC:\n{topic}\n"
        )

        try:
            resp = self.openai_client.responses.create(
                model="gpt-4o-mini",
                tools=[
                    {"type": "file_search", "vector_store_ids": [VECTOR_STORE_ID]},
                    {"type": "web_search"},
                ],
                input=prompt,
            )

            report_text = resp.output_text or "Research ready."
            # Cache the last report for saving
            self._last_report_topic = topic
            self._last_report_content = report_text
            self._last_report_format = format

            # Say a quick one-liner
            first_line = report_text.splitlines()[0] if report_text else "Research ready."
            await context.session.say(first_line[:220])
            return report_text
        except Exception as e:
            logging.exception("deep_research_report failed")
            return f"Deep research error: {e}"


    # ---- Save last report (no huge args; robust & fast) ----
    @function_tool(
        description="Save the most recently generated report to disk as DOCX or PDF (default: docx).")
    async def save_last_report(self, context: RunContext, file_type: str = "docx") -> str:
        if file_type not in ("docx", "pdf"):
            return "Invalid file_type. Use 'docx' or 'pdf'."
        if not self._last_report_content or not self._last_report_topic:
            return "No report in memory. Ask me to generate a brief or report first."

        try:
            out_path = _save_report_to_file(self._last_report_topic, self._last_report_content, file_type)
            self._last_saved_path = out_path
            await context.session.say(f"Saved your report to {out_path}.")
            logging.info("Report saved at %s", out_path)
            return out_path
        except Exception as e:
            logging.exception("Save failed")
            return f"Save failed: {e}"
    

# ----------- Upload report to drive -------------
    @function_tool(
        description="Upload the most recently saved report to Google Drive. Optionally pass a Drive folder ID.")
    async def upload_last_report_to_drive(self, context: RunContext, folder_id: Optional[str] = None) -> str:
        if not self._last_saved_path or not os.path.exists(self._last_saved_path):
            return "No saved report found. Ask me to save the report first."
        try:
            target_folder = folder_id or (GOOGLE_FOLDER_ID or None)
            meta = upload_file_to_drive(self._last_saved_path, target_folder)
            link = meta.get("webViewLink") or f"https://drive.google.com/file/d/{meta.get('id')}/view"
            msg = f"Uploaded to Google Drive: {meta.get('name')} (link: {link})"
            await context.session.say("Uploaded to Google Drive.")
            logging.info(msg)
            return msg
        except Exception as e:
            logging.exception("Drive upload failed")
            return f"Drive upload failed: {e}"


    # --- TOOL: compose a short email draft ---
    @function_tool(
        description="Draft a short, polite email for a given recipient and topic. Do not send; only return a preview.")
    async def compose_email(self, context: RunContext, to_email: str, topic: str, extra_context: Optional[str] = None) -> str:
        """
        Use the OpenAI Responses API to create a concise email (subject + body).
        The draft is stored and returned for confirmation.
        """
        prompt = f"""Write a concise professional email to {to_email}.
                    Topic: {topic}
                    Extra context (optional): {extra_context or "N/A"}
                    Tone: polite, clear, 80-120 words. 
                    Return plain text with a Subject line (one line) and then the email body."""
        try:
            resp = self.openai_client.responses.create(
                model="gpt-4o-mini",
                input=prompt,
            )
            draft_text = resp.output_text or "Subject: (draft)\n\n(draft body)"
            # naive parse: first line "Subject: ..."
            lines = draft_text.splitlines()
            subject = lines[0].replace("Subject:", "").strip() if lines and lines[0].lower().startswith("subject") else "Update request"
            body = "\n".join(lines[1:]).strip() if len(lines) > 1 else draft_text

            self._pending_email = {"to": to_email, "subject": subject, "body": body}
            return f"Draft ready.\nSubject: {subject}\n\n{body}\n\nSay: 'Send the email' to send it."
        except Exception as e:
            return f"Email draft error: {e}"
        
    # ---- TOOL: send the last composed email after user confirms ---
    @function_tool(
        description="Send the last composed email via Gmail. Use only after the user says 'Send'.")
    async def send_email(self, context: RunContext) -> str:

        """Sends the last draft saved in self._pending_email using Gmail API."""
        if not self._pending_email:
            return "There is no composed email to send. Please ask me to compose one first."
        try:
            meta = send_email_via_gmail(
                to_email=self._pending_email["to"],
                subject=self._pending_email["subject"],
                body=self._pending_email["body"],
            )
            mid = meta.get("id", "unknown-id")
            self._pending_email = None
            return f"Email sent. Gmail message id: {mid}."
        except Exception as e:
            return f"Email send error: {e}"


# ---------- LiveKit entry ----------
async def entrypoint(ctx: agents.JobContext):
    await ctx.connect()

    session = AgentSession(
        stt=lk_openai.STT(model="gpt-4o-transcribe"),
        llm=lk_openai.LLM(model="gpt-4o-mini"),
        tts=lk_openai.TTS(
            model="gpt-4o-mini-tts",
            voice="ash",
            instructions=("Speak quickly and clearly; be concise and confident."),
            speed=1.2,
        ),
        vad=silero.VAD.load(),
        turn_detection="vad",
    )

    await session.start(
        room=ctx.room,
        agent=VoiceAssistant(),
        room_input_options=RoomInputOptions(
            noise_cancellation=noise_cancellation.BVC(),
        ),
    )

    await session.generate_reply(
        instructions=("Greet the user. You can search their files, search the web, "
                      "produce lesson briefs or deep research, and save reports as DOCX or PDF.")
    )

if __name__ == "__main__":
    agents.cli.run_app(agents.WorkerOptions(entrypoint_fnc=entrypoint))
